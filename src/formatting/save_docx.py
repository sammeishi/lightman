import os
from models import Task
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


# 保存doc
def save(task: Task):
    saveFile = get_save_file(task)
    if os.path.exists(saveFile):
        print('doc already saved %s' % saveFile)
        return None
    # 重新保存
    doc = Document()
    formattingJson = read_formatting_json(task.formatting_json_file)
    for chapter in formattingJson['chapterList']:
        insert_heading(doc, chapter['title'])
        insert_paragraph(doc, chapter['content'])
    doc.save(saveFile)
    return None


# 插入大标题
def insert_heading(doc, str):
    # 添加标题并获取其段落对象
    heading = doc.add_heading(str, level=1)
    # 操作标题段落的 Run 对象来设置字体和颜色
    run = heading.runs[0]  # 获取标题的第一个 Run 对象
    run.font.name = '宋体'  # 设置字体为宋体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 确保设置中文字体
    run.font.size = Pt(16)  # 设置字号
    run.font.color.rgb = RGBColor(0, 0, 0)

# 插入段落
def insert_paragraph(doc, str):
    # 强制设置运行字体（针对某些环境的字体问题）
    p = doc.add_paragraph()
    run = p.add_run(str)
    run.font.name = '宋体'  # 强制设置运行字体
    run.font.size = Pt(12)  # 字号
    run.font.color.rgb = RGBColor(0, 0, 0)  ## 添加段落测试默认字体

# 获取格式化后的json
def read_formatting_json(formatting_json_file: str):
    # 保存到文本
    with open(formatting_json_file, 'r', encoding='utf-8') as json_file:
        return json.load(json_file)

# 获取保存文件
def get_save_file(task: Task):
    return '%s/formatting.docx' % (task.output_dir)