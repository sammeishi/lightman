import os

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from models import Task


def save(task: Task, chapters: list[dict]):
    """保存为docx"""
    save_path = os.path.join(task.root_path, 'formatting.docx')
    if os.path.exists(save_path):
        print('already saved %s' % save_path)
        return
    # 重新保存
    doc = Document()
    for chapter in chapters:
        insert_heading(doc, chapter['title'])
        insert_paragraph(doc, chapter['content'])
    doc.save(save_path)


def insert_heading(doc, text):
    # 插入大标题
    # 添加标题并获取其段落对象
    heading = doc.add_heading(text, level=1)
    # 操作标题段落的 Run 对象来设置字体和颜色
    run = heading.runs[0]  # 获取标题的第一个 Run 对象
    run.font.name = '宋体'  # 设置字体为宋体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 确保设置中文字体
    run.font.size = Pt(16)  # 设置字号
    run.font.color.rgb = RGBColor(0, 0, 0)


def insert_paragraph(doc, text):
    # 插入段落
    # 强制设置运行字体（针对某些环境的字体问题）
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'  # 强制设置运行字体
    run.font.size = Pt(12)  # 字号
    run.font.color.rgb = RGBColor(0, 0, 0)  ## 添加段落测试默认字体
